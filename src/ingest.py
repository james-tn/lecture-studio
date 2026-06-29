"""Ingestion: turn source material (image, PDF, Word, text) into authoring material.

Two modes:
- **vision** (default for images/PDF): pages are rasterized and read by gpt-5.4's vision so
  layout is preserved — text, LaTeX equations, tables, and *descriptions of diagrams/figures
  with their spatial arrangement*. This captures things plain text extraction misses (geometry
  figures, graphs, boxed callouts, multi-column layout).
- **text**: fast extraction via pypdf / python-docx (no model call).
"""
from __future__ import annotations

import base64
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import azure_client

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}

VISION_PROMPT = (
    "You are a meticulous document transcriber for a math-lecture pipeline. Transcribe this "
    "page FAITHFULLY, preserving its structure and layout. Rules:\n"
    "- Output clean GitHub-flavored Markdown.\n"
    "- Render ALL mathematics as LaTeX (inline $...$ or display $$...$$), exactly as written.\n"
    "- Reproduce tables as Markdown tables; preserve headings, lists, and reading order.\n"
    "- For every diagram, figure, graph, or geometric construction, add a `> [FIGURE] ...` "
    "note that precisely describes it: the shapes, points, labels, axes, curves, arrows, and "
    "their spatial arrangement — enough detail to redraw it.\n"
    "- Note layout cues (columns, boxed/callout content, what is emphasized).\n"
    "Transcribe only what is on the page; do not invent or solve anything."
)


# ---------- text backends ----------
def _pdf_text(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    return "\n\n".join((pg.extract_text() or "") for pg in reader.pages).strip()


def _docx_text(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts).strip()


# ---------- rasterization ----------
def _pdf_to_pngs(path: Path, dpi: int = 150) -> list[bytes]:
    import fitz
    doc = fitz.open(str(path))
    out = []
    for page in doc:
        pix = page.get_pixmap(dpi=dpi)
        out.append(pix.tobytes("png"))
    return out


def _docx_images(path: Path) -> list[bytes]:
    import docx
    d = docx.Document(str(path))
    blobs = []
    for rel in d.part.rels.values():
        if "image" in rel.reltype:
            try:
                blobs.append(rel.target_part.blob)
            except Exception:
                pass
    return blobs


# ---------- vision ----------
def _vision_page(img: bytes, mime: str = "image/png") -> str:
    data_url = f"data:{mime};base64," + base64.b64encode(img).decode()
    messages = [
        {"role": "system", "content": "Transcribe documents precisely; never fabricate."},
        {"role": "user", "content": [
            {"type": "text", "text": VISION_PROMPT},
            {"type": "image_url", "image_url": {"url": data_url}},
        ]},
    ]
    return azure_client.chat(messages, max_completion_tokens=4000)


def _vision_transcribe(images: list[bytes], mime: str = "image/png") -> str:
    if not images:
        return ""
    with ThreadPoolExecutor(max_workers=4) as ex:
        pages = list(ex.map(lambda im: _vision_page(im, mime), images))
    return "\n\n".join(f"## Page {i+1}\n\n{p}" for i, p in enumerate(pages)).strip()


# ---------- entry point ----------
def load_material(path: str | Path, use_vision: bool = True) -> str:
    p = Path(path)
    ext = p.suffix.lower()

    if ext in IMAGE_EXTS:
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else f"image/{ext.lstrip('.')}"
        return _vision_transcribe([p.read_bytes()], mime)

    if ext == ".pdf":
        if use_vision:
            return _vision_transcribe(_pdf_to_pngs(p))
        text = _pdf_text(p)
        if len(text) < 40:
            raise ValueError(f"{p.name}: little extractable text — retry with vision (default).")
        return text

    if ext == ".docx":
        text = _docx_text(p)
        if use_vision:
            fig = _vision_transcribe(_docx_images(p))
            if fig:
                text += "\n\n# Figures (from embedded images)\n\n" + fig
        return text

    if ext in (".txt", ".md", ".markdown", ""):
        return p.read_text()

    raise ValueError(f"unsupported material type: {ext} (use image/.pdf/.docx/.md/.txt)")


def is_visual(path: str | Path) -> bool:
    """True if the material is an image or PDF (can be fed directly to the vision model)."""
    return Path(path).suffix.lower() in (IMAGE_EXTS | {".pdf"})


# OpenAI / Azure GPT-4o-family allow up to 50 images per request; default to a smaller cap.
MAX_IMAGES = 50
HARD_LIMIT = 50


def _vstack(pngs: list[bytes]) -> bytes:
    """Stack PNG pages vertically into a single PNG (shared max width)."""
    from io import BytesIO
    from PIL import Image
    ims = [Image.open(BytesIO(b)).convert("RGB") for b in pngs]
    w = max(i.width for i in ims)
    scaled = [i if i.width == w else i.resize((w, round(i.height * w / i.width))) for i in ims]
    gap = 16
    h = sum(i.height for i in scaled) + gap * (len(scaled) - 1)
    canvas = Image.new("RGB", (w, h), "white")
    y = 0
    for i in scaled:
        canvas.paste(i, (0, y)); y += i.height + gap
    out = BytesIO(); canvas.save(out, format="PNG"); return out.getvalue()


def load_images(path: str | Path, max_images: int = MAX_IMAGES) -> list[str]:
    """Return data-URLs for an image or PDF pages, for direct multimodal authoring.
    If a PDF has more pages than the cap, pages are concatenated (stacked) to fit the cap."""
    import base64
    import math
    cap = max(1, min(max_images, HARD_LIMIT))
    p = Path(path)
    ext = p.suffix.lower()
    if ext in IMAGE_EXTS:
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else f"image/{ext.lstrip('.')}"
        return [f"data:{mime};base64," + base64.b64encode(p.read_bytes()).decode()]
    if ext != ".pdf":
        raise ValueError(f"{p.name}: not an image/PDF — cannot send directly to the model.")

    pages = _pdf_to_pngs(p)
    if len(pages) > cap:
        per = math.ceil(len(pages) / cap)            # pages stacked per combined image
        pages = [_vstack(pages[i:i + per]) for i in range(0, len(pages), per)]
    return ["data:image/png;base64," + base64.b64encode(b).decode() for b in pages]


if __name__ == "__main__":
    import sys
    novis = "--no-vision" in sys.argv
    args = [a for a in sys.argv[1:] if a != "--no-vision"]
    print(load_material(args[0], use_vision=not novis)[:1500])
